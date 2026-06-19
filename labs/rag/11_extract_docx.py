"""
11_extract_docx.py

DOCX 문서에서 제목, 문단, 표 텍스트를 추출한다.

입력:
- enterprise_docs/docx/*.docx

출력:
- extracted_text/docx_extracted.jsonl

실행:
python 11_extract_docx.py
"""

from __future__ import annotations

from docx import Document

from step2_5_config import ENTERPRISE_DOCS_DIR, EXTRACTED_TEXT_DIR, ensure_directories
from step2_5_utils import clean_text, find_files, stable_id, write_jsonl, print_record_summary


def extract_docx_files() -> list[dict]:
    ensure_directories()

    docx_dir = ENTERPRISE_DOCS_DIR / "docx"
    docx_files = find_files(docx_dir, (".docx",))

    records: list[dict] = []

    for docx_path in docx_files:
        try:
            document = Document(str(docx_path))
        except Exception as exc:
            print(f"[DOCX 읽기 실패] {docx_path.name}: {exc}")
            continue

        section_title = ""
        block_index = 0

        for paragraph in document.paragraphs:
            text = clean_text(paragraph.text)
            if not text:
                continue

            style_name = paragraph.style.name if paragraph.style else ""

            if style_name.lower().startswith("heading") or style_name.startswith("제목"):
                section_title = text

            block_index += 1

            records.append(
                {
                    "id": stable_id(docx_path.name, "docx", block_index),
                    "text": text,
                    "metadata": {
                        "file_name": docx_path.name,
                        "document_type": "docx",
                        "block_no": block_index,
                        "section": section_title,
                        "style": style_name,
                        "source_path": str(docx_path),
                        "extractor": "python-docx",
                    },
                }
            )

        # 표 데이터 추출
        for table_index, table in enumerate(document.tables, start=1):
            table_lines: list[str] = []

            for row in table.rows:
                values = [clean_text(cell.text) for cell in row.cells]
                values = [value for value in values if value]

                if values:
                    table_lines.append(" | ".join(values))

            table_text = clean_text("\n".join(table_lines))

            if table_text:
                block_index += 1
                records.append(
                    {
                        "id": stable_id(docx_path.name, "docx", "table", table_index),
                        "text": table_text,
                        "metadata": {
                            "file_name": docx_path.name,
                            "document_type": "docx",
                            "block_no": block_index,
                            "table_no": table_index,
                            "section": section_title,
                            "source_path": str(docx_path),
                            "extractor": "python-docx",
                        },
                    }
                )

    return records


def main() -> None:
    output_path = EXTRACTED_TEXT_DIR / "docx_extracted.jsonl"
    records = extract_docx_files()
    count = write_jsonl(output_path, records)
    print_record_summary("DOCX 텍스트 추출 완료", count, output_path)


if __name__ == "__main__":
    main()
